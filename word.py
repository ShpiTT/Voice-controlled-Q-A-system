# coding=utf-8
"""
Word文档模块
功能：通过大模型生成内容并写入Word文档
"""

import os
import datetime
from openai import OpenAI

# 尝试加载 .env
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# 尝试导入 python-docx
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    print("提示：未安装python-docx，正在安装...")
    os.system("pip install python-docx")
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        DOCX_AVAILABLE = True
    except ImportError:
        print("安装python-docx失败，Word文档功能不可用")
        DOCX_AVAILABLE = False

# API配置（复用LLM模块的配置）
API_KEY = os.getenv("ALI_API_KEY", "")
BASE_URL = os.getenv("ALI_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1")
MODEL = os.getenv("ALI_MODEL", "qwen-flash")

# 文档保存目录
DOCS_DIR = os.path.join(os.path.dirname(__file__), "documents")


def generate_article(topic, article_type="文章"):
    """
    调用大模型生成文章内容
    :param topic: 文章主题，如"保护环境"
    :param article_type: 文章类型，如"文章"、"作文"、"报告"等
    :return: (success, content) - 成功返回(True, 文章内容)，失败返回(False, 错误信息)
    """
    prompt = f"""
请根据以下主题撰写一篇{article_type}：

主题：{topic}

要求：
1. 内容充实，结构完整，包含标题、正文（可分段落）
2. 语言流畅，逻辑清晰
3. 字数在500-1000字左右
4. 适合写入Word文档的格式

请直接输出文章内容，不需要额外的说明。
"""
    
    try:
        client = OpenAI(
            api_key=API_KEY,
            base_url=BASE_URL
        )
        
        print(f"[Word] 正在生成关于「{topic}」的{article_type}...")
        
        completion = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": "你是一位专业的写作助手，擅长撰写各类文章。请直接输出文章内容，格式清晰，适合写入文档。"},
                {"role": "user", "content": prompt}
            ],
            stream=False,
            temperature=0.7
        )
        
        content = completion.choices[0].message.content.strip()
        print(f"[Word] 文章生成成功，共{len(content)}字")
        return True, content
    
    except Exception as e:
        error_msg = f"生成文章失败：{str(e)}"
        print(f"[Word] {error_msg}")
        return False, error_msg


def create_word_document(title, content, filename=None):
    """
    创建Word文档并写入内容
    :param title: 文档标题
    :param content: 文档内容
    :param filename: 文件名（不含扩展名），默认使用标题+时间戳
    :return: (success, filepath/error_msg)
    """
    if not DOCX_AVAILABLE:
        return False, "python-docx未安装，无法创建Word文档"
    
    try:
        # 确保文档目录存在
        if not os.path.exists(DOCS_DIR):
            os.makedirs(DOCS_DIR)
        
        # 生成文件名
        if not filename:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            # 清理标题中的非法字符
            safe_title = "".join(c for c in title if c.isalnum() or c in "_ -")[:20]
            filename = f"{safe_title}_{timestamp}"
        
        filepath = os.path.join(DOCS_DIR, f"{filename}.docx")
        
        # 创建文档
        doc = Document()
        
        # 添加标题
        title_paragraph = doc.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加创建时间
        time_str = datetime.datetime.now().strftime("%Y年%m月%d日 %H:%M")
        time_para = doc.add_paragraph(f"创建时间：{time_str}")
        time_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        time_para.runs[0].font.size = Pt(10)
        
        # 添加分隔线
        doc.add_paragraph("—" * 40)
        
        # 处理内容，按段落添加
        paragraphs = content.split('\n')
        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                # 检查是否是标题行（以#开头或者是短句）
                if para_text.startswith('#'):
                    # Markdown风格的标题
                    level = para_text.count('#', 0, 4)
                    heading_text = para_text.lstrip('#').strip()
                    doc.add_heading(heading_text, level=min(level, 3))
                elif len(para_text) < 30 and not para_text.endswith('。') and not para_text.endswith('：'):
                    # 可能是小标题
                    heading = doc.add_heading(para_text, level=2)
                else:
                    # 正文段落
                    p = doc.add_paragraph()
                    p.paragraph_format.first_line_indent = Inches(0.3)  # 首行缩进
                    run = p.add_run(para_text)
                    run.font.size = Pt(12)
        
        # 保存文档
        doc.save(filepath)
        
        print(f"[Word] 文档已保存：{filepath}")
        return True, filepath
    
    except Exception as e:
        error_msg = f"创建文档失败：{str(e)}"
        print(f"[Word] {error_msg}")
        return False, error_msg


def write_document(topic, article_type="文章"):
    """
    完整流程：生成文章并写入Word文档
    :param topic: 文章主题
    :param article_type: 文章类型
    :return: (success, message)
    """
    # 第一步：调用大模型生成文章
    success, content = generate_article(topic, article_type)
    if not success:
        return False, content
    
    # 第二步：创建Word文档
    title = f"关于{topic}的{article_type}"
    success, result = create_word_document(title, content)
    
    if success:
        return True, f"已成功创建文档：{result}"
    else:
        return False, result


def parse_write_command(command):
    """
    解析写入文档的指令，提取主题和类型
    :param command: 用户指令，如"写入一篇保护环境文章"
    :return: (topic, article_type)
    """
    # 移除常见的前缀词
    command = command.replace("打开文档", "").replace("写入", "").replace("帮我", "")
    command = command.replace("创建", "").replace("生成", "").replace("写", "")
    command = command.strip()
    
    # 提取文章类型
    article_types = ["文章", "作文", "报告", "论文", "总结", "计划", "方案", "心得", "感想", "日记", "故事"]
    article_type = "文章"  # 默认类型
    
    for t in article_types:
        if t in command:
            article_type = t
            command = command.replace(t, "").strip()
            break
    
    # 移除"一篇"、"一份"等量词
    for word in ["一篇", "一份", "一个", "一段"]:
        command = command.replace(word, "")
    
    # 剩余的就是主题
    topic = command.strip()
    
    # 如果主题为空，设置默认值
    if not topic:
        topic = "未指定主题"
    
    return topic, article_type


if __name__ == "__main__":
    # 测试
    print("=" * 50)
    print("Word文档生成测试")
    print("=" * 50)
    
    # 测试指令解析
    test_commands = [
        "写入一篇保护环境文章"
    ]
    
    print("\n指令解析测试：")
    for cmd in test_commands:
        topic, article_type = parse_write_command(cmd)
        print(f"  指令: {cmd}")
        print(f"  主题: {topic}, 类型: {article_type}")
        print()
    
    # 实际生成测试（需要API配置）
    if API_KEY:
        print("\n文档生成测试：")
        success, msg = write_document("保护环境", "文章")
        print(f"结果: {msg}")
    else:
        print("\n未配置API_KEY，跳过文档生成测试")

